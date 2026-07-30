from pathlib import Path
from docx import Document
from docx.shared import Inches
from PIL import Image

out = Path("tmp/fixtures")
out.mkdir(parents=True, exist_ok=True)
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
doc.add_heading("Service Agreement and Fee Schedule", level=1)
doc.add_paragraph("This document contains a table, long clauses, and fee language for upload extraction verification.")
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
for cell, value in zip(table.rows[0].cells, ["Item", "Frequency", "Amount"]):
    cell.text = value
for row in [("Base service", "Monthly", "$89.99"), ("Administrative fee", "One time", "$19.95"), ("Technology fee", "Monthly", "$8.50")]:
    cells = table.add_row().cells
    for cell, value in zip(cells, row):
        cell.text = value
doc.add_heading("Renewal and Cancellation", level=2)
doc.add_paragraph("The agreement renews automatically for twelve months unless written cancellation is received at least sixty days before renewal. A cancellation charge of $125 may apply.")
doc.save(out / "contract-with-table.docx")

source = Image.open("test-invoice.png").convert("RGB")
source.save(out / "phone-photo.jpg", quality=92)
source.save(out / "phone-photo.webp", quality=90)
source.rotate(90, expand=True).save(out / "rotated-phone-photo.jpg", quality=92)
